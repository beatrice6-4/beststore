from django.contrib import admin
from django.http import HttpResponse, HttpResponseRedirect
from django.utils.html import format_html
from django.urls import reverse
import csv
from CDMIS.models import (
    Group, Payment, Activity, Training, Service, Update, Booking, Requirement, Document, FinancialAccount, Withdrawal
)

# ========================= CUSTOM ADMIN STYLING =========================
admin.site.site_header = "🏢 CDMIS Admin Panel"
admin.site.site_title = "CDMIS Administration"
admin.site.index_title = "Welcome to Community Development Management System"


# ========================= GROUP ADMIN =========================
@admin.register(Group)
class GroupAdmin(admin.ModelAdmin):
    """Beautiful Group Admin Interface"""
    list_display = ('colored_group_name', 'members_count', 'status_badge', 'payments_count')
    search_fields = ('name',)
    actions = ['add_payment_action']
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def save_model(self, request, obj, form, change):
        """Save group and redirect to add payment"""
        super().save_model(request, obj, form, change)
        
        # If this is a new group (not changing existing one), redirect to add payment
        if not change:
            # Store group ID in session for the payment form
            request.session['group_id'] = obj.id
            return HttpResponseRedirect(
                reverse('admin:CDMIS_payment_add') + f'?group={obj.id}'
            )
    
    def response_add(self, request, obj, post_url_kwargs=None):
        """Override response after adding group"""
        # Redirect to payment add page
        return HttpResponseRedirect(
            reverse('admin:CDMIS_payment_add') + f'?group={obj.id}'
        )
    
    def colored_group_name(self, obj):
        """Display group name with icon"""
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-users"></i> {}</span>',
            obj.name
        )
    colored_group_name.short_description = '👥 Group Name'
    
    def members_count(self, obj):
        """Display member count as badge"""
        count = getattr(obj, 'member_count', 0)
        return format_html(
            '<span style="background: #667eea; color: #fff; padding: 6px 12px; border-radius: 20px; font-weight: 600; font-size: 12px;">👤 {} Members</span>',
            count
        )
    members_count.short_description = '👥 Members'
    
    def payments_count(self, obj):
        """Display payment count"""
        count = Payment.objects.filter(group=obj).count()
        return format_html(
            '<span style="background: #28a745; color: #fff; padding: 6px 12px; border-radius: 20px; font-weight: 600; font-size: 12px;">💰 {} Payments</span>',
            count
        )
    payments_count.short_description = '💰 Payments'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Active</span>'
        )
    status_badge.short_description = '✓ Status'
    
    def add_payment_action(self, request, queryset):
        """Action to add payment for selected groups"""
        if queryset.count() == 1:
            group = queryset.first()
            return HttpResponseRedirect(
                reverse('admin:CDMIS_payment_add') + f'?group={group.id}'
            )
        self.message_user(request, '❌ Please select only one group to add payment.')
    add_payment_action.short_description = '💰 Add Payment for Selected Group'


# ========================= ACTIVITY ADMIN =========================
@admin.register(Activity)
class ActivityAdmin(admin.ModelAdmin):
    """Beautiful Activity Admin Interface"""
    list_display = ('colored_activity_name', 'category_badge', 'status_badge')
    search_fields = ('name',)
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_activity_name(self, obj):
        """Display activity name with icon"""
        return format_html(
            '<span style="color: #764ba2; font-weight: 600;"><i class="fas fa-tasks"></i> {}</span>',
            obj.name
        )
    colored_activity_name.short_description = '📋 Activity'
    
    def category_badge(self, obj):
        """Display category as badge"""
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">🏷️ Activity</span>'
        )
    category_badge.short_description = '🏷️ Category'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #cce5ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Enabled</span>'
        )
    status_badge.short_description = '✓ Status'


# ========================= TRAINING ADMIN =========================
@admin.register(Training)
class TrainingAdmin(admin.ModelAdmin):
    """Beautiful Training Admin Interface"""
    list_display = ('colored_training_name', 'training_type_badge', 'status_badge')
    search_fields = ('name',)
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_training_name(self, obj):
        """Display training name with icon"""
        return format_html(
            '<span style="color: #28a745; font-weight: 600;"><i class="fas fa-graduation-cap"></i> {}</span>',
            obj.name
        )
    colored_training_name.short_description = '🎓 Training'
    
    def training_type_badge(self, obj):
        """Display training type as badge"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">📚 Training</span>'
        )
    training_type_badge.short_description = '📚 Type'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Active</span>'
        )
    status_badge.short_description = '✓ Status'


# ========================= SERVICE ADMIN =========================
@admin.register(Service)
class ServiceAdmin(admin.ModelAdmin):
    """Beautiful Service Admin Interface"""
    list_display = ('colored_service_name', 'service_type_badge', 'status_badge')
    search_fields = ('name',)
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_service_name(self, obj):
        """Display service name with icon"""
        return format_html(
            '<span style="color: #fd7e14; font-weight: 600;"><i class="fas fa-cog"></i> {}</span>',
            obj.name
        )
    colored_service_name.short_description = '🔧 Service'
    
    def service_type_badge(self, obj):
        """Display service type as badge"""
        return format_html(
            '<span style="background: #fff3cd; color: #856404; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">🏢 Service</span>'
        )
    service_type_badge.short_description = '🏢 Type'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #cfe2ff; color: #084298; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Enabled</span>'
        )
    status_badge.short_description = '✓ Status'



@admin.register(Payment)
class PaymentAdmin(admin.ModelAdmin):
    list_display = ('group', 'amount', 'payment_date', 'notes')
    list_filter = ('group', 'payment_date')
    search_fields = ('group__name', 'notes')
    fieldsets = (
        ('💰 Payment Details', {
            'fields': ('group', 'amount', 'payment_date'),
            'classes': ('wide',),
        }),
        ('📝 Additional Info', {
            'fields': ('notes',),
            'classes': ('wide', 'collapse'),
        }),
    )
    actions = ['download_financial_report_word']

    def download_financial_report_word(self, request, queryset):
        """Download financial report as Word document for selected payments"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        from collections import defaultdict
        
        # Group payments by date
        payments_by_date = defaultdict(list)
        total_amount = 0
        
        for payment in queryset.order_by('payment_date'):
            payments_by_date[payment.payment_date].append(payment)
            total_amount += payment.amount
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('PAYMENT RECORDS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph('Selected Dates Financial Report')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        
        # Add report generation date
        report_date = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%d %B %Y at %H:%M:%S")}')
        report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add space
        
        # Add summary section
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_cells = summary_table.rows[0].cells
        summary_cells[0].text = 'Total Payment Records'
        summary_cells[1].text = str(len(queryset))
        
        summary_cells = summary_table.rows[1].cells
        summary_cells[0].text = 'Number of Dates'
        summary_cells[1].text = str(len(payments_by_date))
        
        summary_cells = summary_table.rows[2].cells
        summary_cells[0].text = 'Number of Groups'
        summary_cells[1].text = str(queryset.values('group').distinct().count())
        
        summary_cells = summary_table.rows[3].cells
        summary_cells[0].text = 'Grand Total Amount'
        summary_cells[1].text = f'Ksh {total_amount:,.2f}'
        
        doc.add_paragraph()  # Add space
        
        # Add detailed payments section
        for payment_date in sorted(payments_by_date.keys()):
            date_heading = doc.add_heading(f'Date: {payment_date.strftime("%d %B %Y (%A)")}', level=2)
            date_heading.runs[0].font.color.rgb = RGBColor(0, 56, 179)  # Blue color
            
            date_payments = payments_by_date[payment_date]
            date_subtotal = sum(p.amount for p in date_payments)
            
            # Create table for payments on this date
            payment_table = doc.add_table(rows=len(date_payments) + 2, cols=3)
            payment_table.style = 'Light Grid Accent 1'
            
            # Add headers
            header_cells = payment_table.rows[0].cells
            header_cells[0].text = 'Group Name'
            header_cells[1].text = 'Amount (Ksh)'
            header_cells[2].text = 'Notes'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add payment rows
            for idx, payment in enumerate(date_payments):
                row_cells = payment_table.rows[idx + 1].cells
                row_cells[0].text = payment.group.name
                row_cells[1].text = f'{payment.amount:,.2f}'
                row_cells[2].text = payment.notes if payment.notes else '—'
            
            # Add subtotal row
            subtotal_cells = payment_table.rows[-1].cells
            subtotal_cells[0].text = 'DATE SUBTOTAL'
            subtotal_cells[1].text = f'{date_subtotal:,.2f}'
            subtotal_cells[2].text = ''
            
            # Style subtotal row
            for cell in subtotal_cells[:2]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            
            doc.add_paragraph()  # Add space between dates
        
        # Add grand total section
        doc.add_paragraph()
        grand_total_heading = doc.add_heading('GRAND TOTAL', level=2)
        grand_total_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)  # Dark red color
        
        grand_total_para = doc.add_paragraph(f'Total Amount: Ksh {total_amount:,.2f}')
        grand_total_para.runs[0].bold = True
        grand_total_para.runs[0].font.size = Pt(14)
        grand_total_para.runs[0].font.color.rgb = RGBColor(178, 34, 34)
        
        # Prepare response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=financial_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(response)
        return response
    
    download_financial_report_word.short_description = '📄 Download as Word (Selected Dates)'

    def get_form(self, request, obj=None, **kwargs):
        form = super().get_form(request, obj, **kwargs)
        if 'group' in request.GET and 'group' in form.base_fields:
            try:
                group_id = request.GET.get('group')
                form.base_fields['group'].initial = group_id
            except Exception:
                pass
        return form
    
    def colored_group(self, obj):
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-users"></i> {}</span>',
            obj.group.name
        )
    colored_group.short_description = '👥 Group'
    
    def colored_amount(self, obj):
        return format_html(
            '<span style="color: #28a745; font-weight: 600; font-size: 14px;">💰 Ksh {}</span>',
            f'{obj.amount:,.2f}'
        )
    colored_amount.short_description = '💰 Amount'
    
    def payment_date_badge(self, obj):
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 11px;">📅 {}</span>',
            obj.payment_date.strftime('%d %b %Y') if obj.payment_date else 'N/A'
        )
    payment_date_badge.short_description = '📅 Date'
    
    def status_badge(self, obj):
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Recorded</span>'
        )
    status_badge.short_description = '✓ Status'

    def download_payments_csv(self, request, queryset):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename=payments.csv'
        writer = csv.writer(response)
        writer.writerow(['Group', 'Amount', 'Date', 'Notes'])  # REMOVED 'Recorded By'
        for payment in queryset:
            writer.writerow([
                payment.group.name,
                payment.amount,
                payment.payment_date,
                payment.notes
            ])
        return response

    download_payments_csv.short_description = "📥 Download selected payments as CSV"

# ========================= REQUIREMENT ADMIN =========================
@admin.register(Requirement)
class RequirementAdmin(admin.ModelAdmin):
    """Beautiful Requirement Admin Interface"""
    list_display = ('colored_requirement_title', 'description_preview', 'status_badge')
    search_fields = ('title', 'description')
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_requirement_title(self, obj):
        """Display requirement title with icon"""
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-checklist"></i> {}</span>',
            obj.title
        )
    colored_requirement_title.short_description = '📋 Requirement'
    
    def description_preview(self, obj):
        """Display description preview"""
        preview = obj.description[:60] + '...' if len(obj.description) > 60 else obj.description
        return format_html(
            '<span style="color: #6c757d; font-style: italic;">"{}"</span>',
            preview
        )
    description_preview.short_description = '📝 Description'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Active</span>'
        )
    status_badge.short_description = '✓ Status'


# ========================= DOCUMENT ADMIN =========================
@admin.register(Document)
class DocumentAdmin(admin.ModelAdmin):
    """Beautiful Document Admin Interface"""
    list_display = ('colored_document_title', 'uploaded_by_badge', 'uploaded_at_badge', 'status_badge')
    search_fields = ('title',)
    list_filter = ('uploaded_at',)
    readonly_fields = ('uploaded_at',)
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_document_title(self, obj):
        """Display document title with icon"""
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-file-pdf"></i> {}</span>',
            obj.title
        )
    colored_document_title.short_description = '📄 Document'
    
    def uploaded_by_badge(self, obj):
        """Display uploader as badge"""
        return format_html(
            '<span style="background: #cfe2ff; color: #084298; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">👤 {}</span>',
            obj.uploaded_by.username if obj.uploaded_by else 'Unknown'
        )
    uploaded_by_badge.short_description = '👤 Uploaded By'
    
    def uploaded_at_badge(self, obj):
        """Display upload date as badge"""
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 11px;">📅 {}</span>',
            obj.uploaded_at.strftime('%d %b %Y') if obj.uploaded_at else 'N/A'
        )
    uploaded_at_badge.short_description = '📅 Date'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Available</span>'
        )
    status_badge.short_description = '✓ Status'


# ========================= UPDATE ADMIN =========================
@admin.register(Update)
class UpdateAdmin(admin.ModelAdmin):
    """Beautiful Update Admin Interface"""
    list_display = ('colored_update_title', 'date_badge', 'created_by_badge', 'status_badge')
    search_fields = ('title', 'content')
    list_filter = ('date', 'created_by')
    readonly_fields = ('date', 'content_display')
    actions = ['download_updates_csv']
    
    fieldsets = (
        ('📢 Update Information', {
            'fields': ('title', 'date', 'created_by'),
            'classes': ('wide',),
        }),
        ('📝 Content', {
            'fields': ('content_display',),
            'classes': ('wide',),
        }),
    )
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_update_title(self, obj):
        """Display update title with icon"""
        return format_html(
            '<span style="color: #764ba2; font-weight: 600;"><i class="fas fa-bullhorn"></i> {}</span>',
            obj.title
        )
    colored_update_title.short_description = '📢 Update'
    
    def date_badge(self, obj):
        """Display date as badge"""
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 11px;">📅 {}</span>',
            obj.date.strftime('%d %b %Y') if obj.date else 'N/A'
        )
    date_badge.short_description = '📅 Date'
    
    def created_by_badge(self, obj):
        """Display creator as badge"""
        return format_html(
            '<span style="background: #cfe2ff; color: #084298; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">👤 {}</span>',
            obj.created_by.username if obj.created_by else 'System'
        )
    created_by_badge.short_description = '👤 Created By'
    
    def status_badge(self, obj):
        """Display status badge"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Published</span>'
        )
    status_badge.short_description = '✓ Status'
    
    def content_display(self, obj):
        """Display full content"""
        return format_html(
            '<div style="background: #f8f9fa; padding: 12px; border-radius: 8px; border-left: 4px solid #667eea;"><p style="margin: 0; color: #2c3e50; line-height: 1.6;">{}</p></div>',
            obj.content
        )
    content_display.short_description = '📝 Full Content'

    def download_updates_csv(self, request, queryset):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename=updates.csv'
        writer = csv.writer(response)
        writer.writerow(['Title', 'Content', 'Date', 'Created By'])
        for update in queryset:
            writer.writerow([
                update.title,
                update.content,
                update.date,
                update.created_by.username if update.created_by else '',
            ])
        return response

    download_updates_csv.short_description = "📥 Download selected updates as CSV"


# ========================= BOOKING ADMIN =========================
@admin.register(Booking)
class BookingAdmin(admin.ModelAdmin):
    """Beautiful Booking Admin Interface"""
    list_display = ('colored_user', 'colored_update', 'booked_at_badge', 'status_badge')
    search_fields = ('user__username', 'update__title')
    list_filter = ('booked_at', 'update')
    readonly_fields = ('booked_at',)
    
    fieldsets = (
        ('📋 Booking Details', {
            'fields': ('user', 'update', 'booked_at'),
            'classes': ('wide',),
        }),
    )
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_user(self, obj):
        """Display user with icon"""
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-user-circle"></i> {}</span>',
            obj.user.username
        )
    colored_user.short_description = '👤 User'
    
    def colored_update(self, obj):
        """Display update title"""
        return format_html(
            '<span style="color: #764ba2; font-weight: 600;"><i class="fas fa-calendar-check"></i> {}</span>',
            obj.update.title
        )
    colored_update.short_description = '📢 Update'
    
    def booked_at_badge(self, obj):
        """Display booking date as badge"""
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 11px;">📅 {}</span>',
            obj.booked_at.strftime('%d %b %Y %H:%M') if obj.booked_at else 'N/A'
        )
    booked_at_badge.short_description = '📅 Booked At'
    
    def status_badge(self, obj):
        """Display booking status"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Confirmed</span>'
        )
    status_badge.short_description = '✓ Status'


# ========================= FINANCIAL ACCOUNT ADMIN =========================
@admin.register(FinancialAccount)
class FinancialAccountAdmin(admin.ModelAdmin):
    """Beautiful Financial Account Admin Interface"""
    list_display = ('colored_user', 'colored_balance', 'phone_badge', 'status_badge')
    search_fields = ('user__username', 'phone_number')
    readonly_fields = ('user',)
    
    fieldsets = (
        ('👤 Account Owner', {
            'fields': ('user',),
            'classes': ('wide',),
        }),
        ('💰 Financial Details', {
            'fields': ('balance', 'phone_number'),
            'classes': ('wide',),
        }),
    )
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_user(self, obj):
        """Display user with icon"""
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-user-circle"></i> {}</span>',
            obj.user.username
        )
    colored_user.short_description = '👤 User'
    
    def colored_balance(self, obj):
        """Display balance with currency symbol"""
        if obj.balance >= 0:
            color = '#28a745'
        else:
            color = '#dc3545'
        return format_html(
            '<span style="color: {}; font-weight: 600; font-size: 14px;">💰 Ksh {}</span>',
            color, f'{obj.balance:,.2f}'
        )
    colored_balance.short_description = '💰 Balance'
    
    def phone_badge(self, obj):
        """Display phone number as badge"""
        return format_html(
            '<span style="background: #cfe2ff; color: #084298; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">📱 {}</span>',
            obj.phone_number
        )
    phone_badge.short_description = '📱 Phone'
    
    def status_badge(self, obj):
        """Display account status"""
        return format_html(
            '<span style="background: #d4edda; color: #155724; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ Active</span>'
        )
    status_badge.short_description = '✓ Status'


# ========================= WITHDRAWAL ADMIN =========================
@admin.register(Withdrawal)
class WithdrawalAdmin(admin.ModelAdmin):
    """Beautiful Withdrawal Admin Interface"""
    list_display = (
        'colored_user',
        'colored_amount',
        'method_badge',
        'status_colored',
        'requested_at_badge'
    )
    search_fields = ('user__username', 'phone_number', 'bank_name', 'account_number')
    list_filter = ('status', 'withdrawal_method', 'requested_at')
    readonly_fields = ('requested_at', 'withdrawal_details')
    
    fieldsets = (
        ('👤 User Information', {
            'fields': ('user', 'phone_number'),
            'classes': ('wide',),
        }),
        ('💰 Withdrawal Details', {
            'fields': ('amount', 'withdrawal_method', 'requested_at'),
            'classes': ('wide',),
        }),
        ('🏦 Bank Details (if applicable)', {
            'fields': ('bank_name', 'bank_branch', 'account_number'),
            'classes': ('wide', 'collapse'),
        }),
        ('📊 Full Details', {
            'fields': ('withdrawal_details',),
            'classes': ('wide',),
        }),
    )
    
    actions = ['mark_pending', 'mark_completed', 'mark_rejected']
    
    class Media:
        css = {
            'all': ('admin/css/cdmis_admin.css',)
        }
    
    def colored_user(self, obj):
        """Display user with icon"""
        return format_html(
            '<span style="color: #667eea; font-weight: 600;"><i class="fas fa-user-circle"></i> {}</span>',
            obj.user.username
        )
    colored_user.short_description = '👤 User'
    
    def colored_amount(self, obj):
        """Display amount with currency symbol"""
        return format_html(
            '<span style="color: #28a745; font-weight: 600; font-size: 14px;">💰 Ksh {}</span>',
            f'{obj.amount:,.2f}'
        )
    colored_amount.short_description = '💰 Amount'
    
    def method_badge(self, obj):
        """Display withdrawal method"""
        method_colors = {
            'mpesa': '#28a745',
            'bank': '#667eea',
            'cash': '#fd7e14',
        }
        color = method_colors.get(obj.withdrawal_method.lower(), '#667eea')
        return format_html(
            '<span style="background: {}; color: #fff; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">💳 {}</span>',
            color, obj.get_withdrawal_method_display()
        )
    method_badge.short_description = '💳 Method'
    
    def status_colored(self, obj):
        """Display status with color coding"""
        status_colors = {
            'pending': '#ffc107',
            'completed': '#28a745',
            'rejected': '#dc3545',
        }
        color = status_colors.get(obj.status.lower(), '#6c757d')
        return format_html(
            '<span style="background: {}; color: #fff; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 12px;">✓ {}</span>',
            color, obj.get_status_display()
        )
    status_colored.short_description = '✓ Status'
    
    def requested_at_badge(self, obj):
        """Display requested date as badge"""
        return format_html(
            '<span style="background: #e7f3ff; color: #004085; padding: 6px 12px; border-radius: 6px; font-weight: 600; font-size: 11px;">📅 {}</span>',
            obj.requested_at.strftime('%d %b %Y') if obj.requested_at else 'N/A'
        )
    requested_at_badge.short_description = '📅 Date'
    
    def withdrawal_details(self, obj):
        """Display full withdrawal details"""
        details = f"""
        <div style="background: #f8f9fa; padding: 15px; border-radius: 8px; border-left: 4px solid #667eea;">
            <p><strong>👤 User:</strong> {obj.user.username}</p>
            <p><strong>💰 Amount:</strong> Ksh {obj.amount:,.2f}</p>
            <p><strong>📱 Phone:</strong> {obj.phone_number}</p>
            <p><strong>💳 Method:</strong> {obj.get_withdrawal_method_display()}</p>
            <p><strong>✓ Status:</strong> {obj.get_status_display()}</p>
            <p><strong>📅 Requested:</strong> {obj.requested_at.strftime('%d %b %Y %H:%M') if obj.requested_at else 'N/A'}</p>
        </div>
        """
        return format_html(details)
    withdrawal_details.short_description = '📊 Full Details'
    
    def mark_pending(self, request, queryset):
        """Mark withdrawals as pending"""
        count = queryset.update(status='pending')
        self.message_user(request, f'⏳ {count} withdrawal(s) marked as pending.')
    mark_pending.short_description = '⏳ Mark as Pending'
    
    def mark_completed(self, request, queryset):
        """Mark withdrawals as completed"""
        count = queryset.update(status='completed')
        self.message_user(request, f'✓ {count} withdrawal(s) marked as completed.')
    mark_completed.short_description = '✓ Mark as Completed'
    
    def mark_rejected(self, request, queryset):
        """Mark withdrawals as rejected"""
        count = queryset.update(status='rejected')
        self.message_user(request, f'✗ {count} withdrawal(s) marked as rejected.')
    mark_rejected.short_description = '✗ Mark as Rejected'