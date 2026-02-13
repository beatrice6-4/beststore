from multiprocessing import context
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy
from django.views.generic import ListView, DetailView, CreateView
from .models import Group, Payment, Activity, Training, Service
from django import forms
from django.db.models import Sum
from django.views import View
from datetime import datetime
from collections import defaultdict
from django.contrib.auth.mixins import UserPassesTestMixin

# --- Forms ---
class GroupForm(forms.ModelForm):
    class Meta:
        model = Group
        fields = ['name', 'registration_date', 'description']

class PaymentForm(forms.ModelForm):
    class Meta:
        model = Payment
        fields = ['group', 'amount', 'payment_date', 'notes']

class ActivityForm(forms.ModelForm):
    class Meta:
        model = Activity
        fields = ['group', 'title', 'activity_date', 'description']

class TrainingForm(forms.ModelForm):
    class Meta:
        model = Training
        fields = ['group', 'topic', 'trainer', 'training_date', 'notes']

class ServiceForm(forms.ModelForm):
    class Meta:
        model = Service
        fields = ['group', 'name', 'service_date', 'description']



from django.views.generic import UpdateView, DeleteView

class GroupUpdateView(UserPassesTestMixin, UpdateView):
    model = Group
    form_class = GroupForm
    template_name = 'CDMIS/group_form.html'
    success_url = reverse_lazy('cdmis:group_list')

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

class GroupDeleteView(UserPassesTestMixin, DeleteView):
    model = Group
    template_name = 'CDMIS/group_confirm_delete.html'
    success_url = reverse_lazy('cdmis:group_list')

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")




# --- Group Views ---
class GroupListView(ListView):
    model = Group
    template_name = 'CDMIS/group_list.html'
    context_object_name = 'groups'
    paginate_by = 9  # Show 9 groups per page

class GroupDetailView(DetailView):
    model = Group
    template_name = 'CDMIS/group_detail.html'
    context_object_name = 'group'

class GroupCreateView(CreateView):
    model = Group
    form_class = GroupForm
    template_name = 'CDMIS/group_form.html'
    success_url = reverse_lazy('cdmis:group_list')

# --- Payment Views ---
from django.shortcuts import render
from django.views.generic import ListView
from .models import Payment
from django.db.models import Sum
from collections import defaultdict
from django.contrib.auth.mixins import UserPassesTestMixin

class PaymentListView(UserPassesTestMixin, ListView):
    model = Payment
    template_name = 'CDMIS/payment_list.html'
    context_object_name = 'payments'

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        all_payments = Payment.objects.all().order_by('payment_date')
        context['all_payments'] = all_payments

        # Calculate totals per date
        date_totals = defaultdict(int)
        for payment in all_payments:
            date_totals[payment.payment_date] += payment.amount
        context['date_totals'] = date_totals

        # Group payments by date for easy template rendering
        payments_by_date = defaultdict(list)
        for payment in all_payments:
            payments_by_date[payment.payment_date].append(payment)
        filtered_payments = {k: v for k, v in payments_by_date.items() if k is not None}
        context['payments_by_date'] = sorted(filtered_payments.items())
        return context

class PaymentCreateView(UserPassesTestMixin, CreateView):
    model = Payment
    form_class = PaymentForm
    template_name = 'CDMIS/payment_form.html'
    success_url = reverse_lazy('cdmis:payment_list')

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

# --- Activity Views ---
class ActivityListView(ListView):
    model = Activity
    template_name = 'CDMIS/activity_list.html'
    context_object_name = 'activities'

class ActivityCreateView(CreateView):
    model = Activity
    form_class = ActivityForm
    template_name = 'CDMIS/activity_form.html'
    success_url = reverse_lazy('cdmis:activity_list')

# --- Training Views ---
class TrainingListView(ListView):
    model = Training
    template_name = 'CDMIS/training_list.html'
    context_object_name = 'trainings'

class TrainingCreateView(CreateView):
    model = Training
    form_class = TrainingForm
    template_name = 'CDMIS/training_form.html'
    success_url = reverse_lazy('cdmis:training_list')

# --- Service Views ---
class ServiceListView(ListView):
    model = Service
    template_name = 'CDMIS/service_list.html'
    context_object_name = 'services'

class ServiceCreateView(CreateView):
    model = Service
    form_class = ServiceForm
    template_name = 'CDMIS/service_form.html'
    success_url = reverse_lazy('cdmis:service_list')




from django import forms
from django.http import HttpResponse

class FinanceDateForm(forms.Form):
    dates = forms.MultipleChoiceField(
        choices=[],
        widget=forms.CheckboxSelectMultiple,
        required=True,
        label="Select Dates"
    )

class FinanceView(UserPassesTestMixin, View):
    template_name = 'CDMIS/finance.html'

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

    def get(self, request, *args, **kwargs):
        from collections import defaultdict
        from .models import Group, Payment
        from datetime import datetime

        payments = Payment.objects.select_related('group').order_by('payment_date')
        finance_data = defaultdict(list)
        date_choices = set()
        
        for payment in payments:
            if payment.payment_date:
                finance_data[payment.payment_date].append({
                    'group': payment.group.name,
                    'amount': payment.amount
                })
                date_choices.add(payment.payment_date)

        finance_list = []
        for date, items in finance_data.items():
            date_total = sum(item['amount'] for item in items)
            finance_list.append({
                'date': date,
                'payments': items,
                'date_total': date_total
            })
        finance_list.sort(key=lambda x: x['date'])

        grand_total = payments.aggregate(total=Sum('amount'))['total'] or 0

        date_choices = sorted([d for d in date_choices if d is not None])
        form = FinanceDateForm()
        form.fields['dates'].choices = [(str(d), d.strftime("%b %d, %Y")) for d in date_choices]

        return render(request, self.template_name, {
            'finance_list': finance_list,
            'grand_total': grand_total,
            'form': form,
            'date_choices_count': len(date_choices),
        })

    def post(self, request, *args, **kwargs):
        from datetime import datetime
        from django.http import HttpResponseForbidden
        import csv
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        payments = Payment.objects.select_related('group').order_by('payment_date')
        
        # Get all available dates
        date_choices = sorted(set(payment.payment_date for payment in payments if payment.payment_date))
        form = FinanceDateForm(request.POST)
        form.fields['dates'].choices = [(str(d), d.strftime("%b %d, %Y")) for d in date_choices]

        if form.is_valid():
            selected_dates_str = form.cleaned_data['dates']
            
            # Convert string dates back to date objects
            try:
                selected_dates = []
                for date_str in selected_dates_str:
                    parsed_date = datetime.strptime(date_str, '%Y-%m-%d').date()
                    selected_dates.append(parsed_date)
            except (ValueError, TypeError):
                return self.get(request)
            
            # Filter payments for selected dates only
            selected_payments = payments.filter(payment_date__in=selected_dates)

            # Group payments by date
            payments_by_date = defaultdict(list)
            total_amount = 0
            
            for payment in selected_payments:
                payments_by_date[payment.payment_date].append(payment)
                total_amount += payment.amount

            # Determine format from request
            download_format = request.POST.get('download_format', 'csv')

            if download_format == 'word':
                # Generate Word document
                doc = Document()
                
                # Add title
                title = doc.add_heading('PAYMENT RECORDS REPORT', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add subtitle with date range
                subtitle = doc.add_paragraph('Selected Dates Financial Report')
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle.runs[0].bold = True
                
                # Add report generation date
                report_date = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%d %B %Y at %H:%M:%S")}')
                report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()  # Add space
                
                # Add summary section
                summary_table = doc.add_table(rows=3, cols=2)
                summary_table.style = 'Light Grid Accent 1'
                
                summary_cells = summary_table.rows[0].cells
                summary_cells[0].text = 'Total Payment Records'
                summary_cells[1].text = str(len(selected_payments))
                
                summary_cells = summary_table.rows[1].cells
                summary_cells[0].text = 'Number of Dates'
                summary_cells[1].text = str(len(payments_by_date))
                
                summary_cells = summary_table.rows[2].cells
                summary_cells[0].text = 'Grand Total Amount'
                summary_cells[1].text = f'Ksh {total_amount:,.2f}'
                
                doc.add_paragraph()  # Add space
                
                # Add detailed payments section
                for payment_date in sorted(payments_by_date.keys()):
                    date_heading = doc.add_heading(f'Date: {payment_date.strftime("%d %B %Y (%A)")}', level=2)
                    date_heading.runs[0].font.color.rgb = RGBColor(0, 56, 179)  # Blue color
                    
                    date_payments = payments_by_date[payment_date]
                    date_subtotal = sum(p.amount for p in date_payments)
                    
                    # Create table for payments on this date
                    payment_table = doc.add_table(rows=len(date_payments) + 2, cols=2)
                    payment_table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    header_cells = payment_table.rows[0].cells
                    header_cells[0].text = 'Group Name'
                    header_cells[1].text = 'Amount (Ksh)'
                    
                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Add payment rows
                    for idx, payment in enumerate(date_payments):
                        row_cells = payment_table.rows[idx + 1].cells
                        row_cells[0].text = payment.group.name
                        row_cells[1].text = f'{payment.amount:,.2f}'
                    
                    # Add subtotal row
                    subtotal_cells = payment_table.rows[-1].cells
                    subtotal_cells[0].text = 'DATE SUBTOTAL'
                    subtotal_cells[1].text = f'{date_subtotal:,.2f}'
                    
                    # Style subtotal row
                    for cell in subtotal_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                    
                    doc.add_paragraph()  # Add space between dates
                
                # Add grand total section
                doc.add_paragraph()
                grand_total_heading = doc.add_heading('GRAND TOTAL', level=2)
                grand_total_heading.runs[0].font.color.rgb = RGBColor(178, 34, 34)  # Dark red color
                
                grand_total_para = doc.add_paragraph(f'Total Amount: Ksh {total_amount:,.2f}')
                grand_total_para.runs[0].bold = True
                grand_total_para.runs[0].font.size = Pt(14)
                grand_total_para.runs[0].font.color.rgb = RGBColor(178, 34, 34)
                
                # Prepare response
                response = HttpResponse(
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = 'attachment; filename=finance_report.docx'
                doc.save(response)
                return response

            else:
                # Generate CSV response (default)
                response = HttpResponse(content_type='text/csv')
                response['Content-Disposition'] = 'attachment; filename=finance_summary.csv'
                
                writer = csv.writer(response)
                writer.writerow(['PAYMENT RECORDS - SELECTED DATES'])
                writer.writerow([])
                writer.writerow(['Date', 'Group', 'Amount'])
                writer.writerow(['-' * 50, '-' * 50, '-' * 50])
                
                # Write payments organized by date
                for payment_date in sorted(payments_by_date.keys()):
                    date_payments = payments_by_date[payment_date]
                    date_subtotal = sum(p.amount for p in date_payments)
                    
                    writer.writerow([payment_date.strftime("%Y-%m-%d"), '', ''])
                    
                    for payment in date_payments:
                        writer.writerow([
                            '',
                            payment.group.name,
                            f"{payment.amount}"
                        ])
                    
                    writer.writerow(['[Subtotal]', '', f"{date_subtotal}"])
                    writer.writerow([])
                
                # Add grand total
                writer.writerow(['-' * 50, '-' * 50, '-' * 50])
                writer.writerow(['TOTAL', '', f"{total_amount}"])
                writer.writerow([])
                writer.writerow([f"Records downloaded on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"])
                
                return response

        # If form is not valid, re-render page
        return self.get(request)




from django.views.generic import ListView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib.auth.mixins import UserPassesTestMixin
from django.shortcuts import redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import get_user_model
User = get_user_model()

class UserListView(UserPassesTestMixin, ListView):
    model = User
    template_name = 'CDMIS/user_list.html'
    context_object_name = 'users'

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['inactive_users'] = User.objects.filter(is_active=False).order_by('-date_joined')
        context['active_users'] = User.objects.filter(is_active=True).order_by('-date_joined')
        return context

class UserUpdateView(UserPassesTestMixin, UpdateView):
    model = User
    fields = ['username', 'first_name', 'last_name', 'email', 'is_active', 'is_staff']
    template_name = 'CDMIS/user_form.html'
    success_url = reverse_lazy('cdmis:user_list')

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

class UserDeleteView(UserPassesTestMixin, DeleteView):
    model = User
    template_name = 'CDMIS/user_confirm_delete.html'
    success_url = reverse_lazy('cdmis:user_list')

    def test_func(self):
        return self.request.user.is_staff or self.request.user.is_superuser

    def handle_no_permission(self):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")

def activate_user(request, pk):
    if not (request.user.is_staff or request.user.is_superuser):
        from django.http import HttpResponseForbidden
        return HttpResponseForbidden("ERROR 404, ONLY ADMINS ARE ALLOWED TO VIEW THIS PAGE.")
    user = get_object_or_404(User, pk=pk)
    user.is_active = True
    user.save()
    messages.success(request, f"User {user.username} activated.")
    return redirect('cdmis:user_list')






from django.http import HttpResponse
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from .models import Payment
import os
from django.conf import settings

def download_payments_pdf_by_date(request, payment_date):
    payments = Payment.objects.filter(payment_date=payment_date)
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="payments_{payment_date}.pdf"'
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    import os
    from django.conf import settings

    p = canvas.Canvas(response, pagesize=A4)
    width, height = A4

    # Add GOV logo
    logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'gov.png')
    if os.path.exists(logo_path):
        p.drawImage(logo_path, inch, height - 1.5*inch, width=1.2*inch, preserveAspectRatio=True, mask='auto')

    p.setFont("Helvetica-Bold", 16)
    p.drawString(2.5*inch, height - 1*inch, f"Payments for {payment_date}")

    # Table headers
    p.setFont("Helvetica-Bold", 11)
    y = height - 2*inch
    p.drawString(inch, y, "Group")
    p.drawString(2.5*inch, y, "Amount")
    p.drawString(3.5*inch, y, "Date")
    p.drawString(5*inch, y, "Notes")
    y -= 0.3*inch

    # Table rows
    p.setFont("Helvetica", 10)
    for payment in payments:
        p.drawString(inch, y, str(payment.group.name))
        p.drawString(2.5*inch, y, f"Ksh. {payment.amount}")
        p.drawString(3.5*inch, y, str(payment.payment_date))
        p.drawString(5*inch, y, str(payment.notes)[:40])  # Truncate notes for layout
        y -= 0.25*inch
        if y < 1*inch:
            p.showPage()
            y = height - 1*inch

    p.showPage()
    p.save()
    return response







def contact_messages(request):
    # Your logic here (fetch messages if you have a model)
    return render(request, 'CDMIS/contact_messages.html')






from django.shortcuts import render, redirect
from .models import Order  # Adjust import to your actual Order model
from .forms import OrderForm  # Create this form for adding orders
from django.contrib.auth.decorators import user_passes_test

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def order_list(request):
    orders = Order.objects.all().order_by('-created_at')
    if request.method == 'POST':
        form = OrderForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('cdmis:order_list')
    else:
        form = OrderForm()
    return render(request, 'CDMIS/order_list.html', {'orders': orders, 'form': form})





from django.shortcuts import render, redirect, get_object_or_404
from .models import Order
from .forms import OrderForm

def order_list(request):
    orders = Order.objects.all()
    if request.method == 'POST':
        form = OrderForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('cdmis:order_list')
    else:
        form = OrderForm()
    return render(request, 'CDMIS/order_list.html', {'orders': orders, 'form': form})

def order_edit(request, pk):
    order = get_object_or_404(Order, pk=pk)
    if request.method == 'POST':
        form = OrderForm(request.POST, instance=order)
        if form.is_valid():
            form.save()
            return redirect('cdmis:order_list')
    else:
        form = OrderForm(instance=order)
    return render(request, 'CDMIS/order_edit.html', {'form': form, 'order': order})

def order_delete(request, pk):
    order = get_object_or_404(Order, pk=pk)
    if request.method == 'POST':
        order.delete()
        return redirect('cdmis:order_list')
    return render(request, 'CDMIS/order_delete.html', {'order': order})

from django.contrib.auth.decorators import login_required

@login_required
def profile(request):
    return render(request, 'CDMIS/profile.html')

from .models import Group


def group_members(request, pk):
    group = get_object_or_404(Group, id=pk)
    members = group.members.all()  # Assuming related_name='members'
    return render(request, 'CDMIS/group_members.html', {'group': group, 'members': members})

from .models import Member

def member_list(request):
    members = Member.objects.all()
    return render(request, 'CDMIS/member_list.html', {'members': members})

# views.py
from django.contrib import messages
from django.shortcuts import render, redirect
from .forms import MemberUploadForm
from .models import Member
import openpyxl
from django.contrib.auth.decorators import user_passes_test

import datetime

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def upload_members(request):
    if request.method == 'POST':
        form = MemberUploadForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES['excel_file']
            wb = openpyxl.load_workbook(excel_file)
            ws = wb.active
            rows = list(ws.iter_rows(min_row=2, values_only=True))  # skip header

            for row in rows:
                first_name = row[0]
                middle_name = row[1]
                id_no = row[2]
                date_of_birth = row[3]
                gender = row[4]
                email = row[5] if len(row) > 5 else ''
                member_role = row[6] if len(row) > 6 else ''
                disability = row[7] if len(row) > 7 else ''

                # Robust date_of_birth handling
                dob = None
                if isinstance(date_of_birth, datetime.date):
                    dob = date_of_birth
                elif isinstance(date_of_birth, str):
                    try:
                        dob = datetime.date.fromisoformat(date_of_birth)
                    except Exception:
                        dob = None
                elif date_of_birth is None:
                    dob = None

                Member.objects.update_or_create(
                    id_no=id_no,
                    defaults={
                        'first_name': first_name,
                        'middle_name': middle_name,
                        'date_of_birth': dob,
                        'gender': gender,
                        'email': email,
                        'member_role': member_role,
                        'disability': disability,
                    }
                )
            messages.success(request, "Members uploaded successfully!")
            return redirect('cdmis:member_list')
    else:
        form = MemberUploadForm()
    return render(request, 'CDMIS/upload_members.html', {'form': form})



from django.shortcuts import render
from django.db.models import Sum
from .models import Group, Payment, Training

def cdmis_reports(request):
    total_financials = Payment.objects.aggregate(total=Sum('amount'))['total'] or 0

    group_payments = (
        Group.objects
        .annotate(total_paid=Sum('payments__amount'))  # <-- corrected here
        .values('name', 'total_paid')
    )
    below_5000 = [g for g in group_payments if (g['total_paid'] or 0) < 5000]
    above_5000 = [g for g in group_payments if (g['total_paid'] or 0) >= 5000]

    total_trainings = Training.objects.count()
    total_groups = Group.objects.count()

    context = {
        'total_financials': total_financials,
        'below_5000_groups': below_5000,
        'above_5000_groups': above_5000,
        'total_groups': total_groups,
        'total_trainings': total_trainings,
    }
    return render(request, 'CDMIS/reports.html', context)



from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import user_passes_test
from django import forms
from .models import Group, Member

# --- Case Management Form ---
class CaseManagementForm(forms.Form):
    CASE_CHOICES = [
        ('change_office_bearers', 'Change of Office Bearers'),
        ('add_member', 'Addition of Group Member'),
        ('exit_member', 'Exit of Group Member'),
        ('correct_member', 'Correction of Member Details'),
    ]
    case_type = forms.ChoiceField(choices=CASE_CHOICES, label="Case Type")
    group = forms.ModelChoiceField(queryset=Group.objects.all(), label="Group")
    member = forms.ModelChoiceField(queryset=Member.objects.all(), required=False, label="Member (if applicable)")
    details = forms.CharField(widget=forms.Textarea, required=False, label="Details / Notes")

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def case_management(request):
    form = CaseManagementForm(request.POST or None)
    message = None

    if request.method == 'POST' and form.is_valid():
        case_type = form.cleaned_data['case_type']
        group = form.cleaned_data['group']
        member = form.cleaned_data.get('member')
        details = form.cleaned_data.get('details')

        if case_type == 'change_office_bearers':
            message = f"Office bearers for group '{group.name}' have been updated."
        elif case_type == 'add_member':
            if member:
                group.members.add(member)  # Add member to group
                group.save()
                member_name = member.first_name
                if hasattr(member, 'middle_name') and member.middle_name:
                    member_name += f" {member.middle_name}"
                message = f"Member '{member_name}' added to group '{group.name}'."
            else:
                message = "Please select a member to add."
        elif case_type == 'exit_member':
            if member:
                group.members.remove(member)
                group.save()
                member_name = member.first_name
                if hasattr(member, 'middle_name') and member.middle_name:
                    member_name += f" {member.middle_name}"
                message = f"Member '{member_name}' exited from group '{group.name}'."
            else:
                message = "Please select a member to exit."
        elif case_type == 'correct_member':
            if member:
                member_name = member.first_name
                if hasattr(member, 'middle_name') and member.middle_name:
                    member_name += f" {member.middle_name}"
                message = f"Details for member '{member_name}' have been corrected."
            else:
                message = "Please select a member to correct."

    return render(request, 'CDMIS/case_management.html', {
        'form': form,
        'message': message,
    })

from django.shortcuts import render, redirect
from .forms import GroupForm

def group_create(request):
    if request.method == 'POST':
        form = GroupForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('cdmis:groups')
    else:
        form = GroupForm()
    return render(request, 'CDMIS/group_form.html', {'form': form})


from django.shortcuts import render, redirect
from django.contrib.auth.decorators import user_passes_test
from django.http import HttpResponse
from .models import Requirement
from .forms import RequirementForm
import csv

def requirements_list(request):
    requirements = Requirement.objects.all()
    return render(request, 'CDMIS/requirements.html', {'requirements': requirements})

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def create_requirement(request):
    if request.method == 'POST':
        form = RequirementForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('cdmis:requirements')
    else:
        form = RequirementForm()
    return render(request, 'CDMIS/create_requirement.html', {'form': form})

from django.http import HttpResponse
from .models import Requirement
from docx import Document

def download_requirements_word(request):
    requirements = Requirement.objects.all()
    document = Document()
    document.add_heading('Group Registration Requirements', 0)

    for req in requirements:
        document.add_heading(req.title, level=1)
        document.add_paragraph(req.description)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=group_requirements.docx'
    document.save(response)
    return response


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Update, Booking
from django.contrib import messages

@login_required
def updates(request):
    updates_list = Update.objects.all().order_by('-date')  # Fetch updates from the database

    if request.method == 'POST':
        update_id = request.POST.get('update_id')
        update = get_object_or_404(Update, id=update_id)

        # Check if the user has already booked this update
        if Booking.objects.filter(user=request.user, update=update).exists():
            messages.error(request, "You have already booked this update.")
        else:
            Booking.objects.create(user=request.user, update=update)
            messages.success(request, "You have successfully booked this update.")

    return render(request, 'CDMIS/updates.html', {'updates': updates_list})

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .models import Document
from .forms import DocumentForm

@login_required
def docs(request):
    documents = Document.objects.all().order_by('-uploaded_at')
    if request.user.is_staff or request.user.is_superuser:
        if request.method == 'POST':
            form = DocumentForm(request.POST, request.FILES)
            if form.is_valid():
                doc = form.save(commit=False)
                doc.uploaded_by = request.user
                doc.save()
                return redirect('cdmis:docs')
        else:
            form = DocumentForm()
    else:
        form = None
    return render(request, 'CDMIS/docs.html', {'documents': documents, 'form': form})



from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings
from decimal import Decimal
from .models import FinancialAccount, Withdrawal

BANK_CHOICES = [
    ('', 'Select Bank'),
    ('KCB', 'KCB Bank'),
    ('Equity', 'Equity Bank'),
    ('Cooperative', 'Cooperative Bank'),
    ('Absa', 'Absa Bank'),
    ('Stanbic', 'Stanbic Bank'),
    ('NCBA', 'NCBA Bank'),
    ('Family', 'Family Bank'),
    ('I&M', 'I&M Bank'),
    ('DTB', 'Diamond Trust Bank'),
    ('Standard Chartered', 'Standard Chartered Bank'),
    # ...add more as needed...
]

@login_required
def withdraw_funds(request):
    # Ensure the FinancialAccount exists for the user
    account, created = FinancialAccount.objects.get_or_create(
        user=request.user,
        defaults={'balance': Decimal('6000.00')}
    )

    if request.method == 'POST':
        amount = request.POST.get('amount')
        withdrawal_method = request.POST.get('withdrawal_method')
        phone_number = request.POST.get('phone_number')
        bank_name = request.POST.get('bank_name')
        bank_branch = request.POST.get('bank_branch')
        account_number = request.POST.get('account_number')
        notes = request.POST.get('notes')

        try:
            amount = Decimal(amount)
        except (ValueError, TypeError):
            messages.error(request, "Invalid amount entered.")
            return redirect('cdmis:withdraw_funds')

        # Validation
        if amount > account.balance:
            messages.error(request, "Insufficient balance.")
        elif amount <= 0:
            messages.error(request, "Withdrawal amount must be greater than zero.")
        elif not withdrawal_method:
            messages.error(request, "Please select a withdrawal method.")
        elif withdrawal_method in ['mpesa', 'cash_pickup'] and not phone_number:
            messages.error(request, "Phone number is required for this method.")
        elif withdrawal_method == 'bank_transfer' and (not bank_name or not bank_branch or not account_number):
            messages.error(request, "Bank name, branch, and account number are required for bank transfer.")
        else:
            # Deduct the amount and create a withdrawal request
            account.balance -= amount
            account.save()

            withdrawal = Withdrawal.objects.create(
                user=request.user,
                amount=amount,
                withdrawal_method=withdrawal_method,
                phone_number=phone_number if withdrawal_method in ['mpesa', 'cash_pickup'] else None,
                bank_name=bank_name if withdrawal_method == 'bank_transfer' else None,
                bank_branch=bank_branch if withdrawal_method == 'bank_transfer' else None,
                account_number=account_number if withdrawal_method == 'bank_transfer' else None,
                notes=notes,
                status='Pending'
            )

            # Send pending notification email
            send_mail(
                subject='Withdrawal Request Pending',
                message=f'Dear {request.user.first_name},\n\nYour withdrawal request of Ksh {withdrawal.amount} is pending and will be processed soon.',
                from_email=settings.DEFAULT_FROM_EMAIL,
                recipient_list=[request.user.email],
                fail_silently=True,
            )

            messages.success(request, f"Withdrawal request for {amount} has been submitted and is pending.")
            return redirect('cdmis:withdrawal_list')

    return render(request, 'CDMIS/withdraw_funds.html', {
        'account': account,
        'bank_choices': BANK_CHOICES
    })



from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import FinancialAccount, Withdrawal
from decimal import Decimal

@login_required
def withdraw_funds(request):
    # Ensure the FinancialAccount exists for the user
    account, created = FinancialAccount.objects.get_or_create(
        user=request.user,
        defaults={'balance': Decimal('6000.00')}  # Default balance is 6,000
    )

    if request.method == 'POST':
        amount = request.POST.get('amount')
        phone_number = request.POST.get('phone_number')

        try:
            # Convert the amount to Decimal
            amount = Decimal(amount)
        except (ValueError, TypeError):
            messages.error(request, "Invalid amount entered.")
            return redirect('cdmis:withdraw_funds')

        # Check if the user has sufficient balance
        if amount > account.balance:
            messages.error(request, "Insufficient balance.")
        elif amount <= 0:
            messages.error(request, "Withdrawal amount must be greater than zero.")
        elif not phone_number:
            messages.error(request, "Phone number is required.")
        else:
            # Deduct the amount and create a withdrawal request
            account.balance -= amount
            account.save()

            Withdrawal.objects.create(user=request.user, amount=amount, phone_number=phone_number, status='Pending')
            messages.success(request, f"Withdrawal request for {amount} has been submitted.")

            # Redirect to the withdrawal list page
            return redirect('cdmis:withdrawal_list')

    return render(request, 'CDMIS/withdraw_funds.html', {'account': account})



from django.views.generic import ListView
from .models import Withdrawal

class WithdrawalListView(ListView):
    model = Withdrawal
    template_name = 'CDMIS/withdrawal_list.html'
    context_object_name = 'withdrawals'

    def get_queryset(self):
        # Show only withdrawals for the logged-in user
        return Withdrawal.objects.filter(user=self.request.user).order_by('-requested_at')





from django.contrib.auth.decorators import user_passes_test
from django.shortcuts import render, redirect, get_object_or_404
from .models import FinancialAccount
from django.contrib import messages
from django import forms

# Form for editing the balance
class EditBalanceForm(forms.ModelForm):
    class Meta:
        model = FinancialAccount
        fields = ['balance']

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def edit_balance(request, pk):
    account = get_object_or_404(FinancialAccount, pk=pk)
    if request.method == 'POST':
        form = EditBalanceForm(request.POST, instance=account)
        if form.is_valid():
            form.save()
            messages.success(request, f"Balance for {account.user.username} has been updated.")
            return redirect('cdmis:financial_accounts')  # Redirect to a list of financial accounts
    else:
        form = EditBalanceForm(instance=account)
    return render(request, 'CDMIS/edit_balance.html', {'form': form, 'account': account})


@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def financial_accounts(request):
    accounts = FinancialAccount.objects.all()
    return render(request, 'CDMIS/financial_accounts.html', {'accounts': accounts})


from django.views.generic import DetailView
from .models import Withdrawal

class WithdrawalDetailView(DetailView):
    model = Withdrawal
    template_name = 'CDMIS/withdrawal_detail.html'
    context_object_name = 'withdrawal'


# ========================= ADMIN PAYMENT REPORT VIEW =========================
from django.contrib.auth.decorators import user_passes_test
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from django.conf import settings

class PaymentDateSelectForm(forms.Form):
    """Form to select payment dates for report generation"""
    dates = forms.MultipleChoiceField(
        choices=[],
        widget=forms.CheckboxSelectMultiple,
        required=True,
        label="Select Payment Dates"
    )

@user_passes_test(lambda u: u.is_staff or u.is_superuser)
def payment_select_dates(request):
    """View to select dates for financial report"""
    # Get all unique payment dates
    all_payments = Payment.objects.all().order_by('-payment_date').values_list('payment_date', flat=True).distinct()
    date_choices = sorted(set([d for d in all_payments if d is not None]))
    
    form = PaymentDateSelectForm()
    form.fields['dates'].choices = [(str(d), d.strftime("%d %B %Y (%A)")) for d in date_choices]
    
    if request.method == 'POST':
        form = PaymentDateSelectForm(request.POST)
        form.fields['dates'].choices = [(str(d), d.strftime("%d %B %Y (%A)")) for d in date_choices]
        
        if form.is_valid():
            selected_dates_str = form.cleaned_data['dates']
            
            # Convert string dates back to date objects
            try:
                selected_dates = []
                for date_str in selected_dates_str:
                    parsed_date = datetime.strptime(date_str, '%Y-%m-%d').date()
                    selected_dates.append(parsed_date)
            except (ValueError, TypeError):
                form.add_error(None, "Invalid date format. Please try again.")
                return render(request, 'CDMIS/admin_payment_select_dates.html', {'form': form})
            
            # Get payments for selected dates
            selected_payments = Payment.objects.filter(payment_date__in=selected_dates).order_by('payment_date')
            
            # Generate Word document
            doc = Document()
            
            # Add logo
            logo_path = os.path.join(settings.BASE_DIR, 'static', 'images', 'gov.png')
            if os.path.exists(logo_path):
                logo_paragraph = doc.add_paragraph()
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(1.2))
            
            # Add title
            title = doc.add_heading('PAYMENT RECORDS REPORT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add subtitle
            subtitle = doc.add_paragraph('STATE DEPARTMENT FOR SOCIAL DEVELOPMENT - CDMIS')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].bold = True
            
            # Add report generation date
            report_date = doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%d %B %Y at %H:%M:%S")}')
            report_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date range info
            if selected_payments:
                min_date = min(p.payment_date for p in selected_payments)
                max_date = max(p.payment_date for p in selected_payments)
                date_range = doc.add_paragraph(f'Date Range: {min_date.strftime("%d %B %Y")} to {max_date.strftime("%d %B %Y")}')
                date_range.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Add space
            
            # Calculate totals
            total_amount = sum(p.amount for p in selected_payments)
            
            # Add summary section
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            
            summary_cells = summary_table.rows[0].cells
            summary_cells[0].text = 'Total Payment Records'
            summary_cells[1].text = str(len(selected_payments))
            
            summary_cells = summary_table.rows[1].cells
            summary_cells[0].text = 'Number of Dates'
            summary_cells[1].text = str(len(selected_dates))
            
            summary_cells = summary_table.rows[2].cells
            summary_cells[0].text = 'Number of Groups'
            summary_cells[1].text = str(selected_payments.values('group').distinct().count())
            
            summary_cells = summary_table.rows[3].cells
            summary_cells[0].text = 'Grand Total Amount'
            summary_cells[1].text = f'Ksh {total_amount:,.2f}'
            
            doc.add_paragraph()  # Add space
            
            # Add all payments in one table
            heading = doc.add_heading('Payment Details - All Selected Dates', level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 56, 179)
            
            # Create single table for all payments
            total_rows = len(selected_payments) + 2
            payment_table = doc.add_table(rows=total_rows, cols=4)
            payment_table.style = 'Light Grid Accent 1'
            
            # Add headers
            header_cells = payment_table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Group Name'
            header_cells[2].text = 'Amount (Ksh)'
            header_cells[3].text = 'Notes'
            
            # Make header bold and colored
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                # Set header background to dark blue
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_elm = parse_xml(r'<w:shd {} w:fill="003B7F"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            # Add payment rows
            row_idx = 1
            for payment in selected_payments:
                row_cells = payment_table.rows[row_idx].cells
                row_cells[0].text = payment.payment_date.strftime('%d %b %Y')
                row_cells[1].text = payment.group.name
                row_cells[2].text = f'{payment.amount:,.2f}'
                row_cells[3].text = payment.notes if payment.notes else '—'
                row_idx += 1
            
            # Add grand total row
            grand_total_cells = payment_table.rows[-1].cells
            grand_total_cells[0].text = 'GRAND TOTAL'
            grand_total_cells[1].text = ''
            grand_total_cells[2].text = f'{total_amount:,.2f}'
            grand_total_cells[3].text = ''
            
            # Style grand total row
            for cell in grand_total_cells[:3]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(178, 34, 34)
                # Set background color
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFEB9C"/>'.format(nsdecls('w')))
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            doc.add_paragraph()  # Add space
            
            # Add final summary
            doc.add_paragraph()
            final_summary = doc.add_paragraph(f'Total Amount: Ksh {total_amount:,.2f}')
            final_summary.runs[0].bold = True
            final_summary.runs[0].font.size = Pt(14)
            final_summary.runs[0].font.color.rgb = RGBColor(178, 34, 34)
            
            # Prepare response
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename=financial_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc.save(response)
            return response
    
    return render(request, 'CDMIS/admin_payment_select_dates.html', {'form': form})